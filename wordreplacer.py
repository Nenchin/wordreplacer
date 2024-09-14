import streamlit as st
from docx import Document
import io

# Function to replace a word in the document
def replace_word_in_doc(doc, target_word, replacement_content):
    for paragraph in doc.paragraphs:
        if target_word in paragraph.text:
            paragraph.text = paragraph.text.replace(target_word, replacement_content)

# Function to read the content of a .docx file as text
def read_docx(file):
    doc = Document(file)
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text + "\n"
    return content

# Function to create a new .docx file with replaced content
def create_new_docx(base_doc, target_word, replacement_content):
    new_doc = Document()
    
    for paragraph in base_doc.paragraphs:
        if target_word in paragraph.text:
            new_paragraph = paragraph.text.replace(target_word, replacement_content)
            new_doc.add_paragraph(new_paragraph)
        else:
            new_doc.add_paragraph(paragraph.text)
    
    # Return a bytes buffer of the new doc
    buffer = io.BytesIO()
    new_doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app
st.title("Word Replacer in .docx")

# File upload for base document
base_file = st.file_uploader("Upload the file with word to be replaced (.docx file)", type="docx")

# File upload for replacement document
replacement_file = st.file_uploader("Upload the replacement (.docx) file", type="docx")

# Text input for word to replace
target_word = st.text_input("Enter the word to replace")

if st.button("Replace"):
    if base_file and replacement_file and target_word:
        # Read the base document
        base_doc = Document(base_file)
        
        # Read the replacement document content
        replacement_content = read_docx(replacement_file)
        
        # Create a new document with the word replaced
        new_doc_buffer = create_new_docx(base_doc, target_word, replacement_content)
        
        # Allow user to download the new document
        st.download_button(
            label="Download the modified .docx",
            data=new_doc_buffer,
            file_name="modified_doc.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("Please upload both files and enter a word to replace.")
